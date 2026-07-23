import docx

def add_header(doc):
    doc.add_paragraph("PANDIT DEENDAYAL")
    doc.add_paragraph("ENERGY UNIVRSITY")
    doc.add_paragraph("SCHOOL OF TECHNOLOGY")
    doc.add_paragraph("")
    doc.add_paragraph("Course: Industrial Training")
    doc.add_paragraph("Course Code: 20TP310")
    doc.add_paragraph("\n")

def create_week1():
    doc = docx.Document()
    add_header(doc)
    doc.add_heading('Talentra (AI Interviewer) Development Progress Report', 0)
    doc.add_heading('Week 1: Core Foundation, Database Architecture & Scrapers', 1)
    
    doc.add_heading('1. Executive Summary', 2)
    doc.add_paragraph('During the first week of development for the Talentra project, the primary objective was to establish a solid, scalable backend architecture capable of handling data processing, database connections, and external API integrations. We initialized the monorepo structure, set up the PostgreSQL database using Prisma ORM, and built robust data extraction modules to parse candidate resumes and scrape GitHub profiles.')
    
    doc.add_heading('2. System Design & File Layout', 2)
    doc.add_paragraph('The backend structure was scaffolded inside the apps/backend/ directory:')
    doc.add_paragraph('- index.ts: The primary Express server entry point that initiates middleware (CORS, body-parser) and REST routes.')
    doc.add_paragraph('- db.ts: Houses the PostgreSQL connection using @prisma/adapter-pg and pg Pool, with intelligent local vs. production SSL resolution.')
    doc.add_paragraph('- scrapers/: Contains the GitHub scraping infrastructure designed to fetch a user\'s repository metadata securely via proxy patterns (testProxy.ts).')
    
    doc.add_heading('3. Core Backend Components Created', 2)
    doc.add_paragraph('A. Database Connection (db.ts)')
    doc.add_paragraph('Implemented a robust database connection leveraging Prisma Client. The setup includes conditional SSL handling to ensure seamless transitions between local development environments and production deployments.')
    doc.add_paragraph('B. Resume Parsing Architecture')
    doc.add_paragraph('Configured multer for handling multipart form data during resume uploads. Integrated pdf-parse to extract raw text streams from uploaded PDF resumes, which is then structured into JSON format using LLM intelligence.')
    doc.add_paragraph('C. Database Schema Definition')
    doc.add_paragraph('Structured the database to handle \'Interview\' entities with state tracking (Pre, InProgress, Done) and a relational \'Message\' model to log chat transcripts between the candidate and the AI.')
    
    doc.save(r'D:\Working Projects\AI_Interviewer\reports\week1_report_detailed.docx')

def create_week2():
    doc = docx.Document()
    add_header(doc)
    doc.add_heading('Talentra (AI Interviewer) Development Progress Report', 0)
    doc.add_heading('Week 2: Real-time Communication & API Architecture', 1)
    
    doc.add_heading('1. Executive Summary', 2)
    doc.add_paragraph('In the second week, development shifted heavily towards establishing real-time communication protocols and the core intelligence routing of the AI Interviewer. We implemented WebRTC signaling mechanisms via Express to connect clients with OpenAI\'s Realtime Voice API and built an extensive multi-LLM fallback architecture to guarantee high availability.')
    
    doc.add_heading('2. AI Integration & Fallback Engine', 2)
    doc.add_paragraph('A. Unified LLM Function (callLLM)')
    doc.add_paragraph('Developed a highly resilient asynchronous `callLLM` function inside index.ts. This function acts as a router that first attempts to use Gemini, falls back to Groq (Llama-3), OpenRouter, and finally attempts local Ollama connections if external APIs face rate limits.')
    doc.add_paragraph('B. Realtime Voice API integration')
    doc.add_paragraph('Created the `/api/v1/session` endpoint to exchange Session Description Protocol (SDP) payloads. This securely negotiates a direct peer-to-peer WebRTC connection with the `gpt-4o-realtime-preview` model, maintaining extremely low latency for conversational voice interactions.')
    
    doc.add_heading('3. REST API Routes Developed', 2)
    doc.add_paragraph('We finalized the core API lifecycle routes:')
    doc.add_paragraph('- POST /api/v1/pre-interview: Ingests the candidate\'s resume, scrapes their GitHub, stores the consolidated metadata in PostgreSQL, and generates an interview ID.')
    doc.add_paragraph('- POST /api/v1/interview/:id/chat: Formulates intelligent technical questions based on the candidate\'s resume and GitHub portfolio, appending messages to the database.')
    
    doc.save(r'D:\Working Projects\AI_Interviewer\reports\week2_report_detailed.docx')

def create_week3():
    doc = docx.Document()
    add_header(doc)
    doc.add_heading('Talentra (AI Interviewer) Development Progress Report', 0)
    doc.add_heading('Week 3: Frontend Interactive UI & Scoring System', 1)
    
    doc.add_heading('1. Executive Summary', 2)
    doc.add_paragraph('During the final week, the focus was entirely on the client-facing experience and post-interview analytics. We built a sophisticated React component (Interview.tsx) for handling WebRTC media streams, embedded real-time audio visualization, built browser-based fallbacks, and developed the backend comprehensive scoring rubric to evaluate candidate performance.')
    
    doc.add_heading('2. The Interview Component (Interview.tsx)', 2)
    doc.add_paragraph('A. WebRTC and Media Recording')
    doc.add_paragraph('Implemented RTCPeerConnection to stream microphone audio to the backend. Additionally, implemented the MediaRecorder API to record the candidate\'s camera and screen (if shared), muxing them into a downloadable WebM/MP4 file automatically upon interview completion.')
    doc.add_paragraph('B. Audio Context Visualization')
    doc.add_paragraph('Integrated the Web Audio API (AudioContext & AnalyserNode). By calculating the Fast Fourier Transform (FFT) frequencies, we linked audio volume to CSS transformations—making the AI avatar and User avatar pulse with glowing halo effects when speaking.')
    doc.add_paragraph('C. Mock Mode Fallback')
    doc.add_paragraph('In scenarios where the WebRTC connection fails, the frontend automatically falls back to an offline "Mock Mode" using the browser\'s native SpeechRecognition and SpeechSynthesis APIs, ensuring the interview can still proceed seamlessly.')
    
    doc.add_heading('3. Evaluation & Reporting Architecture', 2)
    doc.add_paragraph('A. POST /api/v1/interview/:id/evaluate')
    doc.add_paragraph('Engineered a comprehensive post-interview assessment prompt that ingests the full conversation transcript and GitHub metadata. The LLM evaluates the candidate across 5 core factors (GitHub Code Quality, Tech Depth, System Design, CI/CD Testing, and Professionalism), returning a standardized 0-100 score and a detailed markdown feedback report.')
    
    doc.save(r'D:\Working Projects\AI_Interviewer\reports\week3_report_detailed.docx')

if __name__ == "__main__":
    create_week1()
    create_week2()
    create_week3()
